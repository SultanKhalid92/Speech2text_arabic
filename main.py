import os
import yaml
import torch
import librosa
import noisereduce as nr
import soundfile as sf
import yt_dlp
from pyannote.audio import Pipeline
from faster_whisper import WhisperModel
from transformers import AutoTokenizer, AutoModelForSeq2SeqLM
from docx import Document
import re
import nltk

nltk.download("punkt")

# === Load Configuration ===
CONFIG_PATH = os.path.join(os.path.dirname(__file__), "config", "settings.yaml")

with open(CONFIG_PATH, "r", encoding="utf-8") as file:
    config = yaml.safe_load(file)

YOUTUBE_URL = config["youtube_url"]
AUDIO_PATH = config["audio_path"]
CLEAN_AUDIO_PATH = config["clean_audio_path"]
TRANSCRIPT_PATH = config["transcript_path"]
SUMMARY_PATH = config["summary_path"]
DEVICE = torch.device("cuda" if torch.cuda.is_available() else "cpu")
HF_TOKEN = config["hf_token"]
FFMPEG_LOCATION = config["ffmpeg_location"]

# === Functions ===

def extract_audio(youtube_url, output_wav_path="audio.wav"):
    ydl_opts = {
        'format': 'bestaudio/best',
        'outtmpl': 'downloaded_audio.%(ext)s',
        'postprocessors': [{
            'key': 'FFmpegExtractAudio',
            'preferredcodec': 'wav',
            'preferredquality': '192',
        }],
        'postprocessor_args': ['-ar', '16000'],
        'prefer_ffmpeg': True,
        'ffmpeg_location': FFMPEG_LOCATION
    }

    with yt_dlp.YoutubeDL(ydl_opts) as ydl:
        ydl.download([youtube_url])

    if os.path.exists("downloaded_audio.wav"):
        os.rename("downloaded_audio.wav", output_wav_path)
    else:
        raise FileNotFoundError("FFmpeg postprocessing failed to create WAV file.")

def denoise_audio(input_path, output_path):
    y, sr = librosa.load(input_path, sr=16000)
    reduced_noise = nr.reduce_noise(y=y, sr=sr)
    sf.write(output_path, reduced_noise, sr)

def run_diarization(audio_path, hf_token):
    pipeline = Pipeline.from_pretrained(
        "pyannote/speaker-diarization-3.1",
        use_auth_token=hf_token
    )
    pipeline.to(DEVICE)
    return pipeline({'audio': audio_path})

def transcribe_audio(audio_path):
    model = WhisperModel("medium", device="cpu", compute_type="int8")
    segments, _ = model.transcribe(audio_path, language="ar", beam_size=5, vad_filter=True)

    full_text = ""
    all_segments = []

    for segment in segments:
        text = segment.text.strip()
        full_text += text + " "
        all_segments.append({
            "start": segment.start,
            "end": segment.end,
            "text": text
        })

    return {"text": full_text.strip(), "segments": all_segments}

def save_diarized_transcript(transcript, diarization_result, output_path):
    with open(output_path, "w", encoding="utf-8") as f:
        for segment in transcript['segments']:
            start, end, text = segment['start'], segment['end'], segment['text']
            speaker = "Unknown"
            for turn, _, label in diarization_result.itertracks(yield_label=True):
                if turn.start <= start <= turn.end:
                    speaker = label
                    break
            f.write(f"[{start:.2f} - {end:.2f}] {speaker}: {text.strip()}\n")

def summarize_transcript(input_path, output_path_base):
    model_name = "moussaKam/AraBART"
    tokenizer = AutoTokenizer.from_pretrained(model_name)
    model = AutoModelForSeq2SeqLM.from_pretrained(model_name).to(DEVICE)

    with open(input_path, "r", encoding="utf-8") as f:
        transcript = f.read().strip()

    cleaned_transcript = re.sub(r"\[\d+\.\d+\s*-\s*\d+\.\d+\]\s*(SPEAKER_\d+|Unknown):", "", transcript)
    cleaned_transcript = re.sub(r"\s{2,}", " ", cleaned_transcript).strip()

    def generate_summary(prompt_text, max_len=100):
        input_ids = tokenizer.encode(prompt_text, return_tensors="pt", max_length=1024, truncation=True).to(DEVICE)
        summary_ids = model.generate(
            input_ids,
            max_length=max_len,
            min_length=30,
            num_beams=4,
            length_penalty=1.5,
            no_repeat_ngram_size=3,
            early_stopping=True
        )
        return tokenizer.decode(summary_ids.squeeze(), skip_special_tokens=True)

    intro_summary = generate_summary(f"اكتب مقدمة مختصرة للنص التالي:\n{cleaned_transcript}", max_len=80)
    bullet_summary = generate_summary(f"استخرج أهم النقاط باختصار وبجمل قصيرة:\n{cleaned_transcript}", max_len=120)
    conclusion_summary = generate_summary(f"اكتب خلاصة مختصرة للنص:\n{cleaned_transcript}", max_len=80)

    txt_path = output_path_base + ".txt"
    with open(txt_path, "w", encoding="utf-8") as f:
        f.write("### ملخص الفيديو\n\n")
        f.write("🟢 مقدمة:\n" + intro_summary + "\n\n")
        f.write("📌 نقاط رئيسية:\n" + bullet_summary + "\n\n")
        f.write("🔚 خاتمة:\n" + conclusion_summary + "\n")

    doc_path = output_path_base + ".docx"
    doc = Document()
    doc.add_heading("ملخص الفيديو", level=1)
    doc.add_heading("مقدمة", level=2)
    doc.add_paragraph(intro_summary)
    doc.add_heading("نقاط رئيسية", level=2)
    for point in bullet_summary.split("،"):
        doc.add_paragraph(point.strip(), style='List Bullet')
    doc.add_heading("خاتمة", level=2)
    doc.add_paragraph(conclusion_summary)
    doc.save(doc_path)

# === Pipeline Execution ===

if __name__ == "__main__":
    extract_audio(YOUTUBE_URL, AUDIO_PATH)
    denoise_audio(AUDIO_PATH, CLEAN_AUDIO_PATH)
    diarization = run_diarization(CLEAN_AUDIO_PATH, HF_TOKEN)
    transcript = transcribe_audio(CLEAN_AUDIO_PATH)
    save_diarized_transcript(transcript, diarization, TRANSCRIPT_PATH)
    summarize_transcript(TRANSCRIPT_PATH, SUMMARY_PATH)